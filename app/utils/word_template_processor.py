from urllib.parse import quote

from fastapi.responses import StreamingResponse
from io import BytesIO
from docx import Document
from typing import List
import asyncio
import re

from modules.documents.dto import GenerateDocumentFieldDTO
from repositories.s3_repository import S3Object


class WordTemplateProcessor:
    @staticmethod
    async def fill_template(
        s3_object: S3Object, fields: List[GenerateDocumentFieldDTO]
    ) -> StreamingResponse:
        """
        Fill the DOCX template with provided fields and return as StreamingResponse.

        Args:
            s3_object (S3Object): The original DOCX template as S3Object.
            fields (List[GenerateDocumentFieldDTO]): List of field mappings {name: ..., value: ...}

        Returns:
            StreamingResponse: Filled DOCX ready to download
        """

        loop = asyncio.get_event_loop()
        output_stream = await loop.run_in_executor(
            None, WordTemplateProcessor._process_docx, s3_object, fields
        )

        return StreamingResponse(
            output_stream,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": "attachment; filename=template.docx"},
        )

    @staticmethod
    def _process_docx(
        s3_object: S3Object, fields: List[GenerateDocumentFieldDTO]
    ) -> BytesIO:
        template_stream = BytesIO(s3_object.body)
        doc = Document(template_stream)

        # Сначала заменим известные поля
        replacements = {f"<{field.name}>": field.value for field in fields}

        def replace_known_and_unknown(text: str) -> str:
            # Подставляем известные
            for search_text, replace_text in replacements.items():
                text = text.replace(search_text, replace_text)

            # Оставшиеся плейсхолдеры типа <...> заменим на подчёркивания
            pattern = re.compile(r"<[^<>]+>")
            return pattern.sub(lambda m: "_" * len(m.group(0)), text)

        # Заменяем в параграфах
        for paragraph in doc.paragraphs:
            paragraph.text = replace_known_and_unknown(paragraph.text)

        # Заменяем в таблицах
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell.text = replace_known_and_unknown(cell.text)

        output_stream = BytesIO()
        doc.save(output_stream)
        output_stream.seek(0)

        return output_stream

    @staticmethod
    async def generate_docx_response(text: str, filename: str) -> StreamingResponse:
        """
        Асинхронная генерация ответа FastAPI для скачивания docx-файла из текста.

        :param text: Текст, который будет записан в Word-документ.
        :param filename: Название файла для скачивания (с расширением .docx).
        :return: StreamingResponse для FastAPI.
        """

        # Создаем новый документ
        doc = Document()
        doc.add_paragraph(text)

        # Сохраняем документ в оперативную память
        docx_io = BytesIO()
        doc.save(docx_io)
        docx_io.seek(0)

        # Кодируем имя файла для корректного отображения
        quoted_filename = quote(filename)

        # Возвращаем StreamingResponse
        return StreamingResponse(
            docx_io,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={
                "Content-Disposition": f"attachment; filename*=utf-8''{quoted_filename}"
            },
        )

    @staticmethod
    async def replace_placeholders_with_underscores(
        s3_object: S3Object, filename: str = "updated_template.docx"
    ) -> StreamingResponse:
        """
        Заменяет все плейсхолдеры вида <...> в Word-документе на подчёркивания той же длины и возвращает StreamingResponse.

        Args:
            s3_object (S3Object): Исходный DOCX файл в виде S3Object.
            filename (str): Имя итогового файла для скачивания.

        Returns:
            StreamingResponse: DOCX файл с заменой.
        """

        loop = asyncio.get_event_loop()
        output_stream = await loop.run_in_executor(
            None, WordTemplateProcessor._replace_all_placeholders, s3_object
        )

        quoted_filename = quote(filename)

        return StreamingResponse(
            output_stream,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={
                "Content-Disposition": f"attachment; filename*=utf-8''{quoted_filename}"
            },
        )

    @staticmethod
    def _replace_all_placeholders(s3_object: S3Object) -> BytesIO:
        template_stream = BytesIO(s3_object.body)
        doc = Document(template_stream)

        # Регулярка для поиска <...>
        pattern = re.compile(r"<[^<>]+>")

        def replace_in_text(text: str) -> str:
            return pattern.sub(lambda m: "_" * len(m.group(0)), text)

        # Заменяем в параграфах
        for paragraph in doc.paragraphs:
            paragraph.text = replace_in_text(paragraph.text)

        # Заменяем в таблицах
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell.text = replace_in_text(cell.text)

        output_stream = BytesIO()
        doc.save(output_stream)
        output_stream.seek(0)

        return output_stream
